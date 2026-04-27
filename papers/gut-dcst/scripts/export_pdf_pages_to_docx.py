#!/usr/bin/env python3
from __future__ import annotations
import argparse
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

def parse_args():
    p = argparse.ArgumentParser()
    p.add_argument('pdf')
    p.add_argument('docx')
    p.add_argument('--title', required=True)
    p.add_argument('--timestamp', required=True)
    return p.parse_args()

def main():
    args = parse_args()
    pdf = Path(args.pdf).resolve()
    docx_path = Path(args.docx).resolve()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    tmpdir = Path(tempfile.mkdtemp(prefix='docx_pages_'))
    try:
        prefix = tmpdir / 'page'
        subprocess.run([
            '/opt/homebrew/bin/pdftoppm', '-png', str(pdf), str(prefix)
        ], check=True)
        images = sorted(tmpdir.glob('page-*.png'))
        if not images:
            raise RuntimeError(f'No page images created from {pdf}')

        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = Inches(0.4)
        sec.bottom_margin = Inches(0.4)
        sec.left_margin = Inches(0.45)
        sec.right_margin = Inches(0.45)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(args.title)
        run.bold = True
        run.font.size = Pt(16)

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(f'Built {args.timestamp}')
        meta_run.italic = True
        meta_run.font.size = Pt(10)

        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note.add_run('Word handoff exported page-for-page from the current rendered PDF.')
        note_run.font.size = Pt(9)

        doc.add_page_break()

        max_width = Inches(7.55)
        for i, img in enumerate(images):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(img), width=max_width)
            if i != len(images) - 1:
                doc.add_page_break()

        doc.core_properties.title = args.title
        doc.core_properties.subject = 'Timestamped Word export for manuscript review'
        doc.core_properties.comments = f'Built {args.timestamp} from {pdf.name}'
        doc.save(str(docx_path))
        print(docx_path)
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

if __name__ == '__main__':
    main()
